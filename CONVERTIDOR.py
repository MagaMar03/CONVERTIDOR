#!/usr/bin/env python3
"""
╔══════════════════════════════════════════════════════════════════════════════╗
║              PDF OCR CONVERTER PRO V4 - LAYOUT FIEL (ESTILO ABBYY)           ║
╠══════════════════════════════════════════════════════════════════════════════╣
║  • Agrupa texto en BLOQUES coherentes (no fragmenta)                         ║
║  • Preserva posiciones relativas del documento original                      ║
║  • Usa tablas invisibles para elementos lado a lado                          ║
║  • Preprocesamiento avanzado de imagen para mejor OCR                        ║
║  • Detección automática de alineación e indentación                          ║
╚══════════════════════════════════════════════════════════════════════════════╝
"""

import os
import sys
import io
from dataclasses import dataclass
from typing import List, Dict, Tuple
from collections import defaultdict

# ============================================================================
# VERIFICACIÓN DE DEPENDENCIAS
# ============================================================================

DEPENDENCIES_OK = True
MISSING = []

try:
    import fitz  # PyMuPDF
except ImportError:
    DEPENDENCIES_OK = False
    MISSING.append("PyMuPDF (pip install pymupdf)")

try:
    from PIL import Image, ImageEnhance
    import numpy as np
except ImportError:
    DEPENDENCIES_OK = False
    MISSING.append("Pillow + NumPy (pip install pillow numpy)")

try:
    import cv2
except ImportError:
    DEPENDENCIES_OK = False
    MISSING.append("OpenCV (pip install opencv-python)")

try:
    import pytesseract
except ImportError:
    DEPENDENCIES_OK = False
    MISSING.append("pytesseract (pip install pytesseract)")

try:
    from docx import Document
    from docx.shared import Pt, Inches, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    DEPENDENCIES_OK = False
    MISSING.append("python-docx (pip install python-docx)")

# Configurar Tesseract en Windows
if sys.platform == 'win32' and DEPENDENCIES_OK:
    tesseract_paths = [
        r'C:\Program Files\Tesseract-OCR',
        r'C:\Program Files (x86)\Tesseract-OCR',
        os.path.join(os.path.expanduser('~'), 'AppData', 'Local', 'Programs', 'Tesseract-OCR')
    ]
    for path in tesseract_paths:
        if os.path.exists(path):
            os.environ['PATH'] += os.pathsep + path
            pytesseract.pytesseract.tesseract_cmd = os.path.join(path, 'tesseract.exe')
            break


# ============================================================================
# CLASES DE DATOS
# ============================================================================

@dataclass
class TextBlock:
    """Bloque de texto con posición"""
    text: str
    rel_x: float      # Posición X relativa (0-1)
    rel_y: float      # Posición Y relativa (0-1)
    rel_width: float  # Ancho relativo
    font_size: float  # Tamaño de fuente estimado
    

# ============================================================================
# PROCESADOR DE IMÁGENES
# ============================================================================

class ImageProcessor:
    """Preprocesamiento de imágenes para mejorar OCR"""
    
    @staticmethod
    def enhance_for_ocr(img: Image.Image) -> Image.Image:
        """Mejora imagen para OCR"""
        img_array = np.array(img)
        
        # Convertir a escala de grises
        if len(img_array.shape) == 3:
            gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
        else:
            gray = img_array
        
        # Mejorar contraste con CLAHE
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
        enhanced = clahe.apply(gray)
        
        # Eliminar ruido suavemente
        denoised = cv2.fastNlMeansDenoising(enhanced, h=8)
        
        # Binarización Otsu (óptima para documentos)
        _, binary = cv2.threshold(denoised, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        
        return Image.fromarray(binary)


# ============================================================================
# MOTOR OCR CON BLOQUES
# ============================================================================

class BlockOCREngine:
    """Motor OCR que agrupa texto en bloques coherentes"""
    
    def __init__(self, lang: str = 'spa+eng', dpi: int = 300):
        self.lang = lang
        self.dpi = dpi
    
    def extract_blocks(self, image: Image.Image, enhance: bool = True) -> List[TextBlock]:
        """Extrae bloques de texto de la imagen"""
        
        img_w, img_h = image.size
        
        # Preprocesar si es necesario
        if enhance:
            processed = ImageProcessor.enhance_for_ocr(image)
        else:
            processed = image
        
        # OCR con configuración optimizada
        config = r'--oem 3 --psm 3 -c preserve_interword_spaces=1'
        ocr_data = pytesseract.image_to_data(
            processed,
            lang=self.lang,
            output_type=pytesseract.Output.DICT,
            config=config
        )
        
        # Agrupar palabras por BLOQUE
        blocks_data = defaultdict(lambda: {
            "lines": defaultdict(list),
            "bbox": [float('inf'), float('inf'), 0, 0]
        })
        
        n = len(ocr_data['text'])
        for i in range(n):
            text = ocr_data['text'][i].strip()
            conf = int(ocr_data['conf'][i])
            
            if text and conf > 20:
                block_num = ocr_data['block_num'][i]
                line_num = ocr_data['line_num'][i]
                x, y = ocr_data['left'][i], ocr_data['top'][i]
                w, h = ocr_data['width'][i], ocr_data['height'][i]
                
                blocks_data[block_num]["lines"][line_num].append({
                    'text': text, 'x': x, 'y': y, 'w': w, 'h': h
                })
                
                # Actualizar bounding box
                bbox = blocks_data[block_num]["bbox"]
                bbox[0] = min(bbox[0], x)
                bbox[1] = min(bbox[1], y)
                bbox[2] = max(bbox[2], x + w)
                bbox[3] = max(bbox[3], y + h)
        
        # Convertir a TextBlocks
        result = []
        
        for block_num, data in blocks_data.items():
            bbox = data["bbox"]
            if bbox[0] == float('inf'):
                continue
            
            # Construir texto del bloque línea por línea
            lines_text = []
            total_h = 0
            word_count = 0
            
            for line_num in sorted(data["lines"].keys()):
                words = sorted(data["lines"][line_num], key=lambda w: w['x'])
                line = ' '.join(w['text'] for w in words)
                lines_text.append(line)
                
                for w in words:
                    total_h += w['h']
                    word_count += 1
            
            if not lines_text:
                continue
            
            full_text = '\n'.join(lines_text)
            avg_h = total_h / max(word_count, 1)
            
            # Calcular métricas relativas
            rel_x = bbox[0] / img_w
            rel_y = bbox[1] / img_h
            rel_w = (bbox[2] - bbox[0]) / img_w
            font_size = max(9, min(13, avg_h * 72 / self.dpi))
            
            result.append(TextBlock(
                text=full_text,
                rel_x=rel_x,
                rel_y=rel_y,
                rel_width=rel_w,
                font_size=font_size
            ))
        
        # Ordenar por posición (arriba-abajo, izquierda-derecha)
        result.sort(key=lambda b: (b.rel_y, b.rel_x))
        
        return result


# ============================================================================
# GENERADOR DE WORD CON LAYOUT
# ============================================================================

class WordLayoutGenerator:
    """Genera documentos Word preservando el layout"""
    
    def __init__(self):
        self.doc = None
        self.content_width = 7.3  # pulgadas
    
    @staticmethod
    def _make_cell_invisible(cell):
        """Hace los bordes de una celda invisibles"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{name}')
            border.set(qn('w:val'), 'nil')
            tcBorders.append(border)
        tcPr.append(tcBorders)
    
    def create_document(self, pages_blocks: List[List[TextBlock]]) -> Document:
        """Crea documento Word desde bloques de múltiples páginas"""
        
        self.doc = Document()
        
        # Configurar márgenes
        for section in self.doc.sections:
            section.top_margin = Inches(0.6)
            section.bottom_margin = Inches(0.6)
            section.left_margin = Inches(0.6)
            section.right_margin = Inches(0.6)
        
        for page_num, blocks in enumerate(pages_blocks):
            if page_num > 0:
                self.doc.add_page_break()
            
            self._render_page(blocks)
        
        return self.doc
    
    def _render_page(self, blocks: List[TextBlock]):
        """Renderiza una página agrupando bloques en filas"""
        
        if not blocks:
            return
        
        # Agrupar bloques por filas (Y similar = misma fila)
        rows = self._group_into_rows(blocks)
        
        for row in rows:
            if len(row) == 1:
                # Un solo bloque - párrafo simple
                self._add_single_block(row[0])
            else:
                # Múltiples bloques lado a lado - tabla invisible
                self._add_multi_block_row(row)
    
    def _group_into_rows(self, blocks: List[TextBlock], y_threshold: float = 0.025) -> List[List[TextBlock]]:
        """Agrupa bloques en filas basado en posición Y"""
        
        rows = []
        current_row = []
        last_y = -1
        
        for block in blocks:
            if last_y == -1 or abs(block.rel_y - last_y) < y_threshold:
                current_row.append(block)
            else:
                if current_row:
                    rows.append(sorted(current_row, key=lambda b: b.rel_x))
                current_row = [block]
            last_y = block.rel_y
        
        if current_row:
            rows.append(sorted(current_row, key=lambda b: b.rel_x))
        
        return rows
    
    def _add_single_block(self, block: TextBlock):
        """Agrega un bloque individual como párrafo"""
        
        para = self.doc.add_paragraph()
        
        # Determinar alineación
        if block.rel_x > 0.55:
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif 0.28 < block.rel_x < 0.45 and block.rel_width < 0.45:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Aplicar indentación si no está en el margen izquierdo
            if block.rel_x > 0.10:
                indent = block.rel_x * self.content_width * 0.85
                para.paragraph_format.left_indent = Inches(min(indent, 4))
        
        # Agregar texto
        run = para.add_run(block.text)
        run.font.name = 'Arial'
        run.font.size = Pt(block.font_size)
        
        # Espaciado
        para.paragraph_format.space_after = Pt(8)
        para.paragraph_format.space_before = Pt(2)
    
    def _add_multi_block_row(self, row: List[TextBlock]):
        """Agrega múltiples bloques en una fila usando tabla invisible"""
        
        table = self.doc.add_table(rows=1, cols=len(row))
        table.autofit = True
        
        for idx, block in enumerate(row):
            cell = table.rows[0].cells[idx]
            self._make_cell_invisible(cell)
            
            para = cell.paragraphs[0]
            run = para.add_run(block.text)
            run.font.name = 'Arial'
            run.font.size = Pt(block.font_size)
            
            # Última celda a la derecha, demás a la izquierda
            if idx == len(row) - 1:
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Espacio después de la tabla
        self.doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ============================================================================
# CONVERSOR PRINCIPAL
# ============================================================================

class PDFToWordConverter:
    """Conversor principal PDF a Word con layout fiel"""
    
    def __init__(self, lang: str = 'spa+eng', dpi: int = 300):
        self.lang = lang
        self.dpi = dpi
        self.ocr_engine = BlockOCREngine(lang=lang, dpi=dpi)
        self.word_gen = WordLayoutGenerator()
    
    def convert(self, pdf_path: str, output_path: str, 
                enhance: bool = True, 
                progress_cb=None) -> str:
        """
        Convierte PDF a Word
        
        Args:
            pdf_path: Ruta del PDF
            output_path: Ruta de salida .docx
            enhance: Mejorar imagen para OCR
            progress_cb: Callback de progreso (percent, message)
        """
        
        def progress(pct, msg):
            if progress_cb:
                progress_cb(pct, msg)
            else:
                print(f"[{pct:3d}%] {msg}")
        
        progress(0, "Abriendo PDF...")
        
        doc = fitz.open(pdf_path)
        total_pages = len(doc)
        
        progress(5, f"PDF: {total_pages} página(s)")
        
        # Procesar cada página
        pages_blocks = []
        
        zoom = self.dpi / 72.0
        mat = fitz.Matrix(zoom, zoom)
        
        for page_num in range(total_pages):
            pct = 5 + int((page_num / total_pages) * 80)
            progress(pct, f"Procesando página {page_num + 1}/{total_pages}...")
            
            page = doc[page_num]
            pix = page.get_pixmap(matrix=mat)
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            
            blocks = self.ocr_engine.extract_blocks(img, enhance=enhance)
            pages_blocks.append(blocks)
        
        doc.close()
        
        progress(85, "Generando documento Word...")
        
        word_doc = self.word_gen.create_document(pages_blocks)
        
        progress(95, "Guardando archivo...")
        
        word_doc.save(output_path)
        
        progress(100, "¡Completado!")
        
        return output_path


# ============================================================================
# INTERFAZ GRÁFICA (TKINTER)
# ============================================================================

def run_gui():
    """Ejecuta la interfaz gráfica"""
    
    import tkinter as tk
    from tkinter import ttk, filedialog, messagebox
    import threading
    
    class App:
        def __init__(self, root):
            self.root = root
            self.root.title("PDF OCR Converter Pro V4 - Layout Fiel")
            self.root.geometry("800x680")
            self.root.resizable(False, False)
            
            self.pdf_path = None
            self.setup_ui()
            
            if not DEPENDENCIES_OK:
                self.show_missing_deps()
            else:
                self.check_tesseract()
        
        def setup_ui(self):
            # === HEADER ===
            header = tk.Frame(self.root, bg="#1e3a5f", height=110)
            header.pack(fill=tk.X)
            
            title = tk.Label(header, text="PDF OCR Converter Pro V4",
                           font=("Segoe UI", 26, "bold"), fg="white", bg="#1e3a5f")
            title.pack(pady=8)
            
            subtitle = tk.Label(header, 
                text="⚡ Layout Fiel - Bloques Coherentes (estilo ABBYY)",
                font=("Segoe UI", 11), fg="#7eb8da", bg="#1e3a5f")
            subtitle.pack()
            
            features = tk.Label(header,
                text="✓ No fragmenta texto  ✓ Posiciona correctamente  ✓ Tablas invisibles para layout",
                font=("Segoe UI", 9), fg="#a0c4de", bg="#1e3a5f")
            features.pack(pady=5)
            
            # === MAIN ===
            main = tk.Frame(self.root, bg="#f0f4f8")
            main.pack(fill=tk.BOTH, expand=True, padx=20, pady=15)
            
            # Estado
            status_frame = tk.LabelFrame(main, text="Estado", 
                                        font=("Segoe UI", 10, "bold"), bg="#f0f4f8")
            status_frame.pack(fill=tk.X, pady=8)
            
            self.status_label = tk.Label(status_frame, text="Verificando...",
                                        font=("Segoe UI", 9), bg="#f0f4f8", fg="#4a5568")
            self.status_label.pack(pady=8)
            
            # Selección de archivo
            file_frame = tk.LabelFrame(main, text="1. Seleccionar PDF",
                                      font=("Segoe UI", 11, "bold"), bg="#f0f4f8")
            file_frame.pack(fill=tk.X, pady=8)
            
            file_inner = tk.Frame(file_frame, bg="#f0f4f8")
            file_inner.pack(fill=tk.X, padx=10, pady=10)
            
            self.file_label = tk.Label(file_inner, text="No se ha seleccionado archivo",
                                      font=("Segoe UI", 9), bg="#f0f4f8", fg="#718096",
                                      wraplength=520, anchor="w")
            self.file_label.pack(side=tk.LEFT, fill=tk.X, expand=True)
            
            select_btn = tk.Button(file_inner, text="📁 Seleccionar PDF", 
                                  command=self.select_file,
                                  bg="#2b6cb0", fg="white", font=("Segoe UI", 10, "bold"),
                                  cursor="hand2", padx=18, pady=6, relief=tk.FLAT)
            select_btn.pack(side=tk.RIGHT)
            
            # Opciones
            opts_frame = tk.LabelFrame(main, text="2. Opciones",
                                      font=("Segoe UI", 11, "bold"), bg="#f0f4f8")
            opts_frame.pack(fill=tk.X, pady=8)
            
            # Mejora de imagen
            self.enhance_var = tk.BooleanVar(value=True)
            enhance_chk = tk.Checkbutton(opts_frame, 
                text="🔍 Preprocesar imagen para mejor OCR (recomendado)",
                variable=self.enhance_var, font=("Segoe UI", 10),
                bg="#f0f4f8", activebackground="#f0f4f8")
            enhance_chk.pack(anchor=tk.W, padx=10, pady=5)
            
            # DPI e idioma
            row = tk.Frame(opts_frame, bg="#f0f4f8")
            row.pack(fill=tk.X, padx=10, pady=8)
            
            tk.Label(row, text="Resolución (DPI):", font=("Segoe UI", 10), 
                    bg="#f0f4f8").pack(side=tk.LEFT)
            
            self.dpi_var = tk.StringVar(value="300")
            dpi_combo = ttk.Combobox(row, textvariable=self.dpi_var,
                                    values=["200", "300", "350", "400"],
                                    state="readonly", width=8)
            dpi_combo.pack(side=tk.LEFT, padx=(5, 30))
            
            tk.Label(row, text="Idioma OCR:", font=("Segoe UI", 10),
                    bg="#f0f4f8").pack(side=tk.LEFT)
            
            self.lang_var = tk.StringVar(value="spa+eng")
            lang_combo = ttk.Combobox(row, textvariable=self.lang_var,
                                     values=["spa", "eng", "spa+eng", "por", "fra"],
                                     state="readonly", width=12)
            lang_combo.pack(side=tk.LEFT, padx=5)
            
            # Progreso
            self.prog_frame = tk.Frame(main, bg="#f0f4f8")
            self.prog_frame.pack(fill=tk.X, pady=10)
            
            self.prog_label = tk.Label(self.prog_frame, text="",
                                      font=("Segoe UI", 10), bg="#f0f4f8", fg="#2b6cb0")
            self.prog_label.pack()
            
            self.prog_bar = ttk.Progressbar(self.prog_frame, mode='determinate', length=620)
            
            # Botón convertir
            self.convert_btn = tk.Button(main, text="🚀 CONVERTIR A WORD",
                                        command=self.start_convert,
                                        bg="#38a169", fg="white",
                                        font=("Segoe UI", 14, "bold"),
                                        cursor="hand2", padx=40, pady=14, relief=tk.FLAT)
            self.convert_btn.pack(pady=18)
            
            # Info
            info = tk.Label(main,
                text="V4: Agrupa en bloques coherentes como ABBYY - No fragmenta palabras",
                font=("Segoe UI", 9), bg="#f0f4f8", fg="#718096")
            info.pack()
        
        def show_missing_deps(self):
            self.status_label.config(
                text="❌ Faltan dependencias:\n" + "\n".join(MISSING),
                fg="#c53030"
            )
            self.convert_btn.config(state=tk.DISABLED)
        
        def check_tesseract(self):
            try:
                pytesseract.get_tesseract_version()
                self.status_label.config(
                    text="✅ Sistema listo - Tesseract OCR detectado",
                    fg="#276749"
                )
            except:
                self.status_label.config(
                    text="⚠️ Tesseract no encontrado\nDescarga: github.com/UB-Mannheim/tesseract/wiki",
                    fg="#c05621"
                )
        
        def select_file(self):
            path = filedialog.askopenfilename(
                title="Seleccionar PDF",
                filetypes=[("PDF", "*.pdf"), ("Todos", "*.*")]
            )
            if path:
                self.pdf_path = path
                self.file_label.config(text=f"📄 {os.path.basename(path)}", fg="#2d3748")
        
        def start_convert(self):
            if not self.pdf_path:
                messagebox.showwarning("Aviso", "Selecciona un PDF primero")
                return
            
            self.convert_btn.config(state=tk.DISABLED)
            self.prog_bar.pack(fill=tk.X, pady=5)
            self.prog_bar['value'] = 0
            
            t = threading.Thread(target=self.do_convert, daemon=True)
            t.start()
        
        def update_progress(self, pct, msg):
            self.prog_bar['value'] = pct
            self.prog_label.config(text=msg)
            self.root.update_idletasks()
        
        def do_convert(self):
            try:
                dpi = int(self.dpi_var.get())
                lang = self.lang_var.get()
                enhance = self.enhance_var.get()
                
                converter = PDFToWordConverter(lang=lang, dpi=dpi)
                
                base = os.path.splitext(os.path.basename(self.pdf_path))[0]
                out_dir = os.path.dirname(self.pdf_path)
                out_path = os.path.join(out_dir, f"{base}_OCR_V4.docx")
                
                def prog_cb(pct, msg):
                    self.root.after(0, lambda: self.update_progress(pct, msg))
                
                result = converter.convert(self.pdf_path, out_path, 
                                          enhance=enhance, progress_cb=prog_cb)
                
                self.root.after(0, lambda: self.on_success(result, out_dir))
                
            except Exception as e:
                self.root.after(0, lambda: self.on_error(str(e)))
        
        def on_success(self, result, out_dir):
            self.prog_bar.pack_forget()
            self.convert_btn.config(state=tk.NORMAL)
            self.prog_label.config(text="")
            
            messagebox.showinfo("Éxito",
                f"✅ Conversión completada:\n\n{os.path.basename(result)}\n\n"
                f"Guardado en:\n{out_dir}")
            
            if sys.platform == 'win32':
                os.startfile(out_dir)
            else:
                os.system(f'xdg-open "{out_dir}"')
        
        def on_error(self, error):
            self.prog_bar.pack_forget()
            self.convert_btn.config(state=tk.NORMAL)
            self.prog_label.config(text="Error", fg="#c53030")
            messagebox.showerror("Error", f"Error:\n\n{error}")
    
    root = tk.Tk()
    App(root)
    root.mainloop()


# ============================================================================
# CLI
# ============================================================================

def run_cli():
    """Interfaz de línea de comandos"""
    import argparse
    
    parser = argparse.ArgumentParser(
        description='PDF OCR Converter Pro V4 - Layout Fiel'
    )
    parser.add_argument('input', help='PDF de entrada')
    parser.add_argument('-o', '--output', help='Archivo de salida .docx')
    parser.add_argument('-l', '--lang', default='spa+eng', help='Idiomas OCR')
    parser.add_argument('-d', '--dpi', type=int, default=300, help='Resolución DPI')
    parser.add_argument('--no-enhance', action='store_true', help='No preprocesar imagen')
    
    args = parser.parse_args()
    
    if not DEPENDENCIES_OK:
        print("ERROR: Faltan dependencias:")
        for m in MISSING:
            print(f"  - {m}")
        sys.exit(1)
    
    if not os.path.exists(args.input):
        print(f"ERROR: No existe: {args.input}")
        sys.exit(1)
    
    output = args.output or os.path.splitext(args.input)[0] + "_OCR_V4.docx"
    
    print("\n" + "="*60)
    print(" PDF OCR Converter Pro V4 - Layout Fiel")
    print("="*60 + "\n")
    
    converter = PDFToWordConverter(lang=args.lang, dpi=args.dpi)
    
    try:
        result = converter.convert(args.input, output, enhance=not args.no_enhance)
        print(f"\n✅ Creado: {result}")
    except Exception as e:
        print(f"\n❌ Error: {e}")
        sys.exit(1)


# ============================================================================
# MAIN
# ============================================================================

if __name__ == "__main__":
    if len(sys.argv) > 1:
        run_cli()
    else:
        run_gui()